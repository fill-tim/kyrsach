import csv

import matplotlib.pyplot as plt
from django.shortcuts import render
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import Pt
from docx2pdf import convert
import pythoncom
import math


def add_error(index, mass_date, mass_value, mass_error):
    """ Добавление ошибок в массив"""
    elem = []
    elem.append(index + 2)
    elem.append(mass_date)
    elem.append(mass_value)
    mass_error.append(elem)


def validation_file(mass_date, mass_value):
    """ Проверка файла на буквы в файле"""
    mass_error = []
    for index in range(len(mass_date)):
        try:
            num_date = int(mass_date[index])
            num_value = int(mass_value[index])
            if num_date <= 0 or num_value <= 0:
                add_error(index, mass_date[index], mass_value[index], mass_error)
        except:
            add_error(index, mass_date[index], mass_value[index], mass_error)
    return mass_error


def read(param):
    """ Метод читает файл и возвращает массив со значениями из файла
    """
    mass_date = []
    mass_value = []
    mass_error = []
    accept = True
    try:
        with open(param, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile, delimiter=' ', quotechar='|')
            for row in reader:
                for item in row:
                    obj = item.split(';')
                    mass_date.append(obj[0])
                    mass_value.append(obj[1])
    except:
        accept = False
        return mass_date, mass_value, mass_error, accept

    mass_date = mass_date[1:]
    mass_value = mass_value[1:]
    if len(mass_date) >= 5:
        mass_error = validation_file(mass_date, mass_value)
        if len(mass_error) == 0:
            mass_date = [int(item) for item in mass_date]
            mass_value = [int(item) for item in mass_value]
    else:
        accept = False

    return mass_date, mass_value, mass_error, accept


def get_absolut_chains_gains(mas):
    """ Возвращает массив с цепными значениями абсолютного прироста
    """
    res = []
    for i in range(len(mas) - 1):
        num = mas[i + 1] - mas[i]
        res.append(num)

    return res


def get_absolut_basic_gains(mas):
    """ Возвращает массив с базисными значениями абсолютного прироста
    """
    res = []

    for i in range(1, len(mas)):
        num = mas[i] - mas[0]
        res.append(num)

    return res


def get_chain_growth_rate(mas):
    """
    Нахождение темпа роста.
    Возвращает цепные коэффиценты и проценты
    """
    factor = []
    percent = []

    for i in range(1, len(mas)):
        num = mas[i] / mas[i - 1]
        num_percent = num * 100

        factor.append(round(num, 3))
        percent.append(round(num_percent, 3))

    return factor, percent


def get_basic_growth_rate(mas):
    """
    Нахождение темпа роста.
    Возвращает двумерный массив базисный коэффиценты И проценты
    """
    factor = []
    percent = []

    for i in range(1, len(mas)):
        num = mas[i] / mas[0]
        num_percent = num * 100

        factor.append(round(num, 3))
        percent.append(round(num_percent, 3))

    return factor, percent


def get_gain_chain_rate(mas):
    """
    Нахождение темпа прироста.
    Возвращает массив с процентами ежегодного темпа роста
    """
    res = []

    for i in range(1, len(mas)):
        num = (mas[i] - mas[i - 1]) / mas[i - 1]
        res.append(round(num, 3))

    return res


def get_gain_basic_rate(mas):
    """
    Нахождение темпа прироста.
    Возвращает массив с процентами темпа роста к начальному
    """
    res = []

    for i in range(1, len(mas)):
        num = (mas[i] - mas[0]) / mas[0]
        res.append(round(num, 3))

    return res


def merge_mas(absolut_mas, rate_mas):
    """
    Объединение массива с данными подсчета абсолютных значений прироста
     и массива с данными подсчета ежегодных темпов прироста в двумерный массив
    :param absolut_mas:
    :param rate_mas:
    :return:
    """

    res = []

    for i in range(len(absolut_mas)):
        elem = []
        elem.append(absolut_mas[i])
        elem.append(rate_mas[i])

        res.append(elem)

    return res


def get_absolute_value_percent_gain(mas):
    """
    Нахождение абсолютного значения 1% прироста.
    :param absolut_mas: Массив с данными подсчета абсолютных значений прироста
    :param rate_mas: Массив с данными подсчета ежегодных темпов прироста
    :return:
    """
    res = []

    for i in range(1, len(mas)):
        num = mas[i-1] * 0.01
        res.append(round(num, 0))
    print(res)
    return res


def get_relative_acceleration(mas):
    """
    Вычисление относительного ускорения
    :param mas:
    :return:
    """
    res = [0]

    for i in range(1, len(mas)):
        num = mas[i] - mas[i - 1]
        res.append(round(num, 3))
    return res


def get_lead_factor(mas):
    """
    Вычисление коэффицента опережения
    :param mas:
    :return:
    """
    res = [0]

    for i in range(1, len(mas)):
        num = mas[i] / mas[i - 1]
        res.append(round(num, 3))

    return res


def forecasting_absolut_gain(mas_gain, mas, num):
    """
    Прогнозирование значений на следующие года методом среднего абсолютного прироста
    :param mas:
    :param mas_gain:
    :param num:
    :return:
    """
    date = [item for item in mas[0]]
    value = [item for item in mas[1]]
    avg_absolut_gain = sum(mas_gain) / len(mas_gain)
    if num != '':

        for i in range(1, num + 1):
            item_value = int(value[len(value) - 1] + avg_absolut_gain * i)
            item_date = date[len(date) - 1] + 1

            date.append(item_date)
            value.append(item_value)

    res = [(e1, e2) for e1, e2 in zip(date, value)]

    return res, date, value, avg_absolut_gain


def forecasting_avg_growth_rate(chain_rate, mas, num):
    """ Прогнозирование данных методом среднего значения прироста"""
    date = [item for item in mas[0]]
    value = [item for item in mas[1]]
    avg_growth_rate = math.prod(chain_rate) ** (1 / len(chain_rate))
    if num != '':

        for i in range(1, num + 1):
            item_date = date[len(date) - 1] + 1
            item_value = int(value[len(value) - 1] * avg_growth_rate)

            date.append(item_date)
            value.append(item_value)

    res = [(e1, e2) for e1, e2 in zip(date, value)]

    return res, date, value, avg_growth_rate


def data_output(request):
    """
    Перезаписывает данные с файла, который пользователь загрузил, в существующий файл для дальнейшей работы.
    :param request:
    :return:
    """
    if request.method == 'POST':
        num = request.POST['num']
        # try:
        file = request.FILES['file']
        if num != '':
            num = int(num)
        with open('timeseries/static/timeseries/files/uploaded_file.csv', 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)

        read_file = read('timeseries/static/timeseries/files/uploaded_file.csv')
        if read_file[3]:
            if len(read_file[2]) == 0:

                data = main(read_file, num)
                create_file(data, data['mass'], data['forecast'], data['forecast_date'], data['forecast_value'],
                            data['label'], data['data'],
                            data['forecast_avg_growth'], data['forecast_avg_growth_date'],
                            data['forecast_avg_growth_value'],
                            num)

                return render(request, 'timeseries/data_output.html', context={'data': data, 'num': num})

            else:
                error_message = 'В файле были найдены ошибки!'
                return render(request, 'timeseries/form_add_file.html',
                              {'error_message': error_message, 'error_mass': read_file[2]})
        else:
            error_message = 'Файл должен содержать минимум 5 значений!'
            return render(request, 'timeseries/form_add_file.html', {'error_message': error_message})
        # except:
        #     error_message = 'Файл не был загружен!'
        #     return render(request, 'timeseries/form_add_file.html', {'error_message': error_message})

    else:
        error = ''
        return render(request, 'timeseries/form_add_file.html', {'error': error})


def insert(mas0, mas1, result, basic, chain_rate0, chain_rate1,
           basic_rate0, basic_rate1, gain_chain_rate, gain_basic_rate,
           absolute_value_percent_gain, relative_acceleration, lead_factor):
    """
    Добавляем пустоту в массивы, чтобы уровнять их длины и создаем один общий массив, чтобы передать на страницу
    """
    result.insert(0, 0)
    basic.insert(0, 0)
    chain_rate0.insert(0, 0)
    chain_rate1.insert(0, 0)
    basic_rate0.insert(0, 0)
    basic_rate1.insert(0, 0)
    absolute_value_percent_gain.insert(0, 0)
    gain_chain_rate.insert(0, 0)
    gain_basic_rate.insert(0, 0)
    lead_factor.insert(0, 0)
    relative_acceleration.insert(0, 0)

    new_mass = [[e1, e2, e3, e4, e5, e6, e7, e8, e9, e10, e11, e12, e13] for
                e1, e2, e3, e4, e5, e6, e7, e8, e9, e10, e11, e12, e13 in
                zip(mas0, mas1, result, basic, chain_rate0, chain_rate1,
                    basic_rate0, basic_rate1, gain_chain_rate, gain_basic_rate,
                    absolute_value_percent_gain, relative_acceleration, lead_factor
                    )]

    return new_mass


def main(file, num):
    """ Запускаем все методы для подсчетов """
    mas = file
    result = get_absolut_chains_gains(mas[1])
    basic = get_absolut_basic_gains(mas[1])
    chain_rate = get_chain_growth_rate(mas[1])
    basic_rate = get_basic_growth_rate(mas[1])
    gain_chain_rate = get_gain_chain_rate(mas[1])
    gain_basic_rate = get_gain_basic_rate(mas[1])
    absolute_value_percent_gain = get_absolute_value_percent_gain(mas[1])
    relative_acceleration = get_relative_acceleration(gain_chain_rate)
    lead_factor = get_lead_factor(chain_rate[0])
    forecasting_absolut = forecasting_absolut_gain(result, mas, num)
    avg_growth_rate = forecasting_avg_growth_rate(chain_rate[0], mas, num)

    new_mass = insert(mas[0], mas[1], result, basic, chain_rate[0], chain_rate[1],
                      basic_rate[0], basic_rate[1], gain_chain_rate, gain_basic_rate,
                      absolute_value_percent_gain, relative_acceleration, lead_factor)

    data = {
        'label': mas[0],
        'data': mas[1],
        'mass': new_mass,
        'num': num,
        'forecast': forecasting_absolut[0],
        'forecast_date': forecasting_absolut[1],
        'forecast_value': forecasting_absolut[2],
        'forecast_avg_growth': avg_growth_rate[0],
        'forecast_avg_growth_date': avg_growth_rate[1],
        'forecast_avg_growth_value': avg_growth_rate[2],
        'avg_absolut_gain': forecasting_absolut[3],
        'avg_growth_rate': avg_growth_rate[3],
    }

    return data


def create_table(data, lst_header_name, document):
    """ Создаем таблицу в файле docx формата """
    doc = document
    table = doc.add_table(1, len(data[0]))

    table.style = 'Table Grid'
    table.alignment = WD_TAB_ALIGNMENT.CENTER
    table.autofit = True

    head_cells = table.rows[0].cells

    for i, item in enumerate(lst_header_name):
        p = head_cells[i].paragraphs[0]
        pc = p.add_run(item)
        pc.font.name = 'Times New Roman'
        pc.font.bold = False
        pc.font.size = Pt(10)

        p.alignment = WD_TAB_ALIGNMENT.CENTER

    for row in data:

        cells = table.add_row().cells
        for i, item in enumerate(row):
            cells[i].text = str(item)
            row_item = cells[i].paragraphs[0].runs[0]
            row_item.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_item.font.name = 'Times New Roman'
            row_item.font.bold = False
            row_item.font.size = Pt(9)


def create_diagram(date, value, doc):
    """ Создаем диаграмму в файле docx формата """
    fig, ax = plt.subplots()
    ax.plot(date, value, color='red')
    ax.bar(date, value)

    ax.set_xlabel('Года')
    ax.set_ylabel('Количество(млн)')

    plt.savefig('timeseries/static/timeseries/photo/plot.png')

    doc.add_picture('timeseries/static/timeseries/photo/plot.png')


def create_diagram_forecast(date1, value1, date2, value2, doc):
    """ Создаем диаграмму в файле docx формата """
    fig, ax = plt.subplots()
    ax.plot(date1, value1, color='red')
    ax.bar(date2, value2)

    ax.set_xlabel('Года')
    ax.set_ylabel('Количество(млн)')

    plt.savefig('timeseries/static/timeseries/photo/plot.png')

    doc.add_picture('timeseries/static/timeseries/photo/plot.png')


def add_description(doc):
    """ Добавление в файл текста с описанием ... """
    par = doc.add_paragraph("""
    Δц - Абсолютные приросты (цепные)
    Δб - Абсолютные приросты (базисные)
    Tрц - Темпы роста цепные (коэффициенты)
    Tрц % - Темпы роста цепные (проценты)
    Трб - Темпы роста базисные (коэффициенты)
    Трб % - Темпы роста базисные (проценты)
    Тпц % - Темпы прироста цепные (проценты)
    Тпб % - Темпы прироста базисные (проценты)
    A - Абсолютное значение 1% прироста
    Δ % - Относительное ускорение (проценты)
    Копер - Коэффициент опережения""")
    font = par.runs[0].font
    font.name = 'Times New Roman'
    font.size = Pt(14)


def font_options(par):
    font = par.runs[0].font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    paragraph_format = par.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_table_diagram_forecast(value_first_forecast, value_second_forecast, doc, forecast, forecast_date,
                               forecast_value, forecast_avg_move,
                               forecast_avg_move_date, forecast_avg_move_value, lst_name_st):
    """ Создание таблицы и диаграммы для прогноза """
    par = doc.add_paragraph('Прогноз данных')
    font = par.runs[0].font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    paragraph_format = par.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par = doc.add_paragraph(f'Прогноз по среднему абсолютному приросту равному -> {int(value_first_forecast)}')
    font_options(par)

    create_table(forecast, lst_name_st, doc)
    par = doc.add_paragraph(f'Прогноз по среднему коэффициенту роста равному -> {round(value_second_forecast, 3)}')
    font_options(par)

    create_table(forecast_avg_move, lst_name_st, doc)
    create_diagram_forecast(forecast_date, forecast_value, forecast_avg_move_date, forecast_avg_move_value, doc)
    par = doc.add_paragraph("""
            Синий цвет - прогноз по среднему коэффициенту роста
            Красный цвет - прогноз по среднему абсолютному приросту""")
    font = par.runs[0].font
    font.name = 'Times New Roman'
    font.size = Pt(14)


def create_file(data, mass, forecast, forecast_date, forecast_value,
                date, value, forecast_avg_growth,
                forecast_avg_growth_date, forecast_avg_growth_value, num):
    """ Создаем файл в docx формате"""
    lst_name_ft = ['Год', 'Исходные значения', 'Δц', 'Δб', 'Tрц', 'Tрц %',
                   'Трб', 'Трб %', 'Тпц %', 'Тпб %', 'A', 'Δ %', 'Копер']
    lst_name_st = ['Год', 'Значения']

    doc = Document()
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    par = doc.add_paragraph('Анализ данных')
    font = par.runs[0].font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    paragraph_format = par.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    create_table(mass, lst_name_ft, doc)
    add_description(doc)
    create_diagram(date, value, doc)
    if num != '':
        add_table_diagram_forecast(data['avg_absolut_gain'], data['avg_growth_rate'], doc, forecast, forecast_date,
                                   forecast_value, forecast_avg_growth,
                                   forecast_avg_growth_date, forecast_avg_growth_value, lst_name_st)

    doc.save('timeseries/static/timeseries/files/test.docx')

    convert_pdf()


def convert_pdf():
    """ Конвертирует docx файл в pdf """
    pythoncom.CoInitializeEx(0)
    convert('timeseries/static/timeseries/files/test.docx')


def home(request):
    """ Главная страница
    """
    return render(request, 'timeseries/home.html')


def add_file(request):
    """ Страница с формой добавления файла"""
    return render(request, 'timeseries/form_add_file.html')


def info(request):
    """ Страница с информацией """
    return render(request, 'timeseries/info.html')


def guide(request):
    """ Страница с руководством """
    return render(request, 'timeseries/guide_page.html')
